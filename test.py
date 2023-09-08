import openpyxl
from docx import Document

# Load the Excel file
excel_file_path = input("Enter Your Excl File path")
workbook = openpyxl.load_workbook(excel_file_path)
sheet = workbook.active 

# Create a Word document
doc = Document()

# Iterate through each row in the Excel sheet and create a Word page for each row
for row in sheet.iter_rows(min_row=2, values_only=True):
    page = doc.add_page_break()
    doc.add_paragraph('Row Data:')
    doc.add_paragraph('\n'.join(map(str, row)))

# Save the Word document
word_file_path = input("Enter Your Word File Path")
doc.save(word_file_path)

print(f'Word file "{word_file_path}" created successfully.')